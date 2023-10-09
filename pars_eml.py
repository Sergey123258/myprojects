import email
import io
import sys
import PyPDF2
from docx import Document
import xml.etree.ElementTree as ET

def decode_email_header(header):
    decoded_parts = []
    for part, encoding in email.header.decode_header(header):
        if encoding is None:
            if isinstance(part, bytes):
                decoded_parts.append(part.decode('utf-8'))
            else:
                decoded_parts.append(part)
        else:
            decoded_parts.append(part.decode(encoding))
    return ' '.join(decoded_parts)


def parse_eml_file(eml_file_path):
    try:
        # Открываем и разбираем EML-файл
        with open(eml_file_path, 'rb') as eml_file:
            msg = email.message_from_binary_file(eml_file)  # Используем email.message_from_binary_file для байтового чтения

        # Получаем основные поля EML
        sender = msg.get('From')
        subject_header = msg.get('Subject')
        subject = decode_email_header(subject_header)
        date = msg.get('Date')
        body = ""

        # Получаем тело EML-сообщения
        if msg.is_multipart():
            for part in msg.walk():
                content_type = part.get_content_type()
                content_disposition = str(part.get("Content-Disposition"))

                if "attachment" not in content_disposition:
                    payload = part.get_payload(decode=True)
                    if payload is not None:
                        body = payload.decode('utf-8')  # Используем кодировку utf-8 для текстовых данных

                    # Получаем текстовое содержимое из HTML
                    if content_type == "text/html":
                        try:
                            import html2text
                            converter = html2text.HTML2Text()
                            converter.ignore_links = True
                            body = converter.handle(body)
                        except ImportError:
                            pass
        else:
            payload = msg.get_payload(decode=True)
            if payload is not None:
                body = payload.decode('utf-8')  # Используем кодировку utf-8 для текстовых данных

        # Получаем вложения
        attachments = []
        for part in msg.walk():
            if part.get_content_maintype() == "multipart":
                continue
            filename = part.get_filename()
            if filename:
                content_type = part.get_content_type()
                payload = part.get_payload(decode=True)
                if payload is not None:
                    attachments.append((filename, content_type, payload))

        # Выводим результаты
        print(f"Отправитель: {sender}")
        print(f"Тема: {subject}")
        print(f"Дата: {date}")
        print("Тело сообщения:")
        print(body)

        for filename, content_type, payload in attachments:
            print("\nВложения:")
            print(f"Имя файла: {filename}")
            print(f"Тип файла: {content_type}")

            if content_type == "application/pdf":
                pdf_text = extract_text_from_pdf(payload)
                print("Содержимое PDF:")
                print(pdf_text)
            elif content_type in ["application/msword", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
                doc_text = extract_text_from_doc(payload)
                print("Содержимое DOC/DOCX:")
                print(doc_text)
            elif content_type == "text/xml":
                xml_text = extract_text_from_xml(payload)
                print("Содержимое XML:")
                print(xml_text)
            elif content_type == "application/octet-stream":
                file_type = identify_file_type(payload)
                if file_type == "application/msword":
                    doc_text = extract_text_from_doc(payload)
                    print("Содержимое DOC:")
                    print(doc_text)
                elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                    doc_text = extract_text_from_doc(payload)
                    print("Содержимое DOCX:")
                    print(doc_text)
                else:
                    print("Не поддерживается разбор этого типа вложения.")
            else:
                print("Не поддерживается разбор этого типа вложения.")

    except Exception as e:
        print(f"Ошибка: {str(e)}")

def extract_text_from_pdf(pdf_content):
    pdf_text = ""
    try:
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(pdf_content))
        for page in pdf_reader.pages:
            pdf_text += page.extract_text()
    except Exception as e:
        pdf_text = f"Ошибка при разборе PDF: {str(e)}"
    return pdf_text

def extract_text_from_doc(doc_content):
    doc_text = ""
    try:
        if doc_content.startswith(b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1'):
            doc = Document(io.BytesIO(doc_content))
            for paragraph in doc.paragraphs:
                doc_text += paragraph.text + "\n"
        elif doc_content.startswith(b'PK\x03\x04'):
            docx = Document(io.BytesIO(doc_content))
            for paragraph in docx.paragraphs:
                doc_text += paragraph.text + "\n"
        else:
            doc_text = "Неподдерживаемый формат документа"
    except Exception as e:
        doc_text = f"Ошибка при разборе DOC/DOCX: {str(e)}"
    return doc_text

def extract_text_from_xml(xml_content):
    xml_text = ""
    try:
        root = ET.fromstring(xml_content)
        xml_text = ET.tostring(root, encoding='utf-8').decode()
    except Exception as e:
        xml_text = f"Ошибка при разборе XML: {str(e)}"
    return xml_text

def identify_file_type(data):
    # Простой метод идентификации типа файла на основе его сигнатуры
    if data.startswith(b'\xD0\xCF\x11\xE0\xA1\xB1\x1A\xE1'):
        return "application/msword"
    elif data.startswith(b'PK\x03\x04'):
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    else:
        return None

# Запуск
if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("Использование: python parse_eml.py [путь_к_файлу.eml]")
    else:
        eml_file_path = sys.argv[1]
        parse_eml_file(eml_file_path)
